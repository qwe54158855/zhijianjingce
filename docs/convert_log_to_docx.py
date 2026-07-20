#!/usr/bin/env python3
"""将 DEVELOPMENT_LOG.md 转换为排版精美的 Word 文档"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_FILE = r"D:\cy\docs\DEVELOPMENT_LOG.md"
OUTPUT_FILE = r"D:\cy\docs\DEVELOPMENT_LOG.docx"

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Styles
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.line_spacing = 1.3

    # Heading styles
    for level, (size, color_hex) in {
        1: (18, '0D47A1'),  # 深蓝
        2: (14, '1565C0'),
        3: (12, '1976D2'),
    }.items():
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '微软雅黑'
        hs.font.size = Pt(size)
        hs.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    current_h1 = ""
    current_h2 = ""
    in_table = False
    table_rows_buffer = []
    table_headers = None

    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    def flush_table():
        nonlocal table_headers, table_rows_buffer, in_table
        if not table_headers or not table_rows_buffer:
            table_headers = None
            table_rows_buffer = []
            in_table = False
            return

        # Determine column widths based on content
        num_cols = len(table_headers)
        col_widths = [Cm(14.5 / num_cols)] * num_cols

        table = doc.add_table(rows=1 + len(table_rows_buffer), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        for i, header in enumerate(table_headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(header.strip())
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '1565C0')

        # Data rows
        for ri, row_data in enumerate(table_rows_buffer):
            for ci, cell_text in enumerate(row_data):
                if ci >= num_cols:
                    break
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text.strip())
                run.font.size = Pt(9)
                if ri % 2 == 1:
                    set_cell_shading(cell, 'F0F4F8')

        # Add spacing after table
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_before = Pt(6)
        p_after.paragraph_format.space_after = Pt(6)

        table_headers = None
        table_rows_buffer = []
        in_table = False

    def is_table_row(line):
        stripped = line.strip()
        return stripped.startswith('|') and stripped.endswith('|') and len(stripped) > 3

    def parse_table_row(line):
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        return cells

    def is_separator_row(line):
        return bool(re.match(r'^\|[\s\-:]+\|$', line.strip()))

    def add_para(text, bold=False, italic=False, size=10.5, color=None, space_before=0, space_after=4, indent_level=0):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if indent_level > 0:
            p.paragraph_format.left_indent = Cm(indent_level * 0.8)
        return p

    def parse_list_item(line):
        """Check if line is a list item and return (level, text)"""
        stripped = line.lstrip()
        if stripped.startswith('- '):
            indent = len(line) - len(stripped)
            return indent // 2, stripped[2:]
        if stripped.startswith('1.') or stripped.startswith('2.') or stripped.startswith('3.'):
            indent = len(line) - len(stripped)
            return indent // 2, stripped[3:].strip()
        if stripped.startswith('* '):
            indent = len(line) - len(stripped)
            return indent // 2, stripped[2:]
        return None

    def parse_inline_formatting(text):
        """Remove markdown formatting and emoji, return plain text"""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'^\d+\.\s*', '', text)
        # Remove emoji
        emoji_pattern = re.compile("["
            u"\U0001F600-\U0001F64F"  # emoticons
            u"\U0001F300-\U0001F5FF"  # symbols & pictographs
            u"\U0001F680-\U0001F6FF"  # transport
            u"\U0001F1E0-\U0001F1FF"  # flags
            u"\U00002702-\U000027B0"  # dingbats
            u"\U000024C2-\U0001F251"  # misc
            u"\U0001F900-\U0001F9FF"  # Supplemental Symbols
            u"\U0001FA00-\U0001FA6F"  # Chess Symbols
            u"\U0001FA70-\U0001FAFF"  # Symbols Extended-A
            u"\U00002600-\U000026FF"  # misc symbols
            u"\U00002700-\U000027BF"  # Dingbats
            u"\U0000FE00-\U0000FE0F"  # Variation Selectors
            u"\U0000200D"            # Zero Width Joiner
            u"\U0000200B"            # Zero Width Space
            u"\U0000FE0F"            # Variation Selector-16
        "]+", flags=re.UNICODE)
        text = emoji_pattern.sub('', text)
        return text

    # Process lines
    for line in lines:
        stripped = line.strip()

        # Skip empty lines between sections, but keep some spacing
        if not stripped:
            if in_table:
                continue  # skip empty lines in tables
            continue

        # Headers
        if stripped.startswith('## '):
            flush_table()
            text = stripped[3:].strip()
            # Remove markdown links from headers
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
            doc.add_heading(text, level=2)
            current_h2 = text
            continue

        if stripped.startswith('# '):
            flush_table()
            text = stripped[2:].strip()
            # Title on first page
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_p.add_run(text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
            title_p.paragraph_format.space_before = Pt(60)
            title_p.paragraph_format.space_after = Pt(6)

            # Add subtitle
            sub_p = doc.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub_p.add_run('完整开发日志 · 2026-07-02 至 2026-07-20')
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            doc.add_paragraph()  # spacing
            continue

        if stripped.startswith('### '):
            flush_table()
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            continue

        # Table handling
        if is_table_row(line):
            cells = parse_table_row(line)
            if is_separator_row(line):
                continue  # skip the |---|---| separator row
            if table_headers is None:
                table_headers = cells
            else:
                table_rows_buffer.append(cells)
            in_table = True
            continue

        # If we were in a table and no longer see table content, flush
        if in_table and not is_table_row(line):
            flush_table()

        # Separator
        if stripped.startswith('---') and len(stripped) >= 3:
            # Horizontal rule
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            continue

        # List items
        list_info = parse_list_item(line)
        if list_info:
            level, text = list_info
            text = parse_inline_formatting(text)
            # Check if it's a numbered item like "1." or bold item
            add_para(text, size=10, indent_level=level + 1)
            continue

        # Code blocks
        if stripped.startswith('```'):
            continue

        # Regular paragraphs
        # Remove markdown formatting
        text = parse_inline_formatting(stripped)
        if text:
            # Check for bold-italic patterns indicating key data
            has_key_data = '**' in stripped
            add_para(text, size=10.5)

    flush_table()  # flush any remaining table

    # Add footer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 文档自动生成于 ' + '2026-07-20 —')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.save(OUTPUT_FILE)
    print(f"OK - Output: {OUTPUT_FILE}")
    print(f"    Size: {os.path.getsize(OUTPUT_FILE) / 1024:.0f} KB")
    print(f"    Paragraphs: {len(doc.paragraphs)}")
    print(f"    Tables: {len(doc.tables)}")

if __name__ == '__main__':
    create_document()
